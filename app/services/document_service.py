import io
import logging
from pathlib import Path
from typing import List, Optional

import cloudinary
import cloudinary.uploader
from fastapi import status
from jinja2 import Environment, FileSystemLoader, select_autoescape
from sqlalchemy.orm import Session

from app.models.admins import Admin
from app.models.ai_generations import AiGeneration
from app.models.documents import Document
from app.models.enums import AdminRole, DocumentType
from app.models.submissions import Submission
from app.schema.ai import StructuredCvData
from app.utils.custom_response import error_response, success_response
from app.utils.settings import settings

logger = logging.getLogger(__name__)

# ── Configure Cloudinary ────────────────────────────────────────────────────
# Initialise here so document_service is self-contained regardless of import order.
cloudinary.config(
    cloud_name=settings.CLOUDINARY_CLOUD_NAME,
    api_key=settings.CLOUDINARY_API_KEY,
    api_secret=settings.CLOUDINARY_API_SECRET,
)

logger = logging.getLogger(__name__)

# ── Template engine setup ───────────────────────────────────────────────────
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

_jinja_env = Environment(
    loader=FileSystemLoader(str(TEMPLATES_DIR)),
    autoescape=select_autoescape(["html"]),
)


# ── HTML Rendering ──────────────────────────────────────────────────────────

def render_cv_to_html(cv_data: StructuredCvData, template_name: Optional[str] = None) -> str:
    """
    Render StructuredCvData → HTML string using a Jinja2 template.
    Falls back to 'cv_default.html' if no custom template is provided or found.
    """
    chosen = template_name or "cv_default.html"

    # Graceful fallback: if a custom template name is given but doesn't exist, use default
    try:
        template = _jinja_env.get_template(chosen)
    except Exception:
        logger.warning(f"Template '{chosen}' not found — falling back to cv_default.html")
        template = _jinja_env.get_template("cv_default.html")

    return template.render(cv=cv_data)


# ── PDF Rendering ───────────────────────────────────────────────────────────

def render_pdf_bytes(html: str) -> bytes:
    """Convert an HTML string into PDF bytes using WeasyPrint."""
    try:
        from weasyprint import HTML as WeasyprintHTML
        pdf_bytes = WeasyprintHTML(string=html).write_pdf()
        return pdf_bytes
    except Exception as e:
        logger.error(f"WeasyPrint PDF rendering failed: {e}")
        raise RuntimeError(f"PDF rendering failed: {e}")


# ── DOCX Rendering ──────────────────────────────────────────────────────────

def render_docx_bytes(cv_data: StructuredCvData) -> bytes:
    """
    Build a polished DOCX document programmatically from StructuredCvData.
    Uses python-docx for full structural control.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise RuntimeError(f"python-docx is not installed: {e}")

    doc = DocxDocument()

    # ── Page margins ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── Helper functions ───────────────────────────────────────────────────
    def add_heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        if level == 1:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif level == 2:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            run.font.letter_spacing = Pt(1.5)
            # Underline paragraph bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'E2E8F0')
            pBdr.append(bottom)
            pPr.append(pBdr)
        return p

    def add_two_col(left_text: str, right_text: str, left_bold: bool = False):
        """Add a paragraph with text aligned left and right (tab stop trick)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run_l = p.add_run(left_text)
        run_l.bold = left_bold
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run("\t")
        run_r = p.add_run(right_text)
        run_r.font.size = Pt(9)
        run_r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        # Right-align the tab stop
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9360')  # 6.5 inches in twips
        tabs.append(tab)
        pPr.append(tabs)
        return p

    def add_sub(text: str, color: tuple = (0x25, 0x63, 0xEB)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(*color)
        run.bold = True
        return p

    def add_body(text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        return p

    def add_bullet(text: str):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        return p

    def add_spacer():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)

    # ── Personal Info Block ────────────────────────────────────────────────
    info = cv_data.personal_info
    add_heading(info.full_name, level=1)
    if info.target_role:
        p = doc.add_paragraph()
        run = p.add_run(info.target_role.upper())
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run.font.letter_spacing = Pt(1)

    contact_parts = []
    if info.email:
        contact_parts.append(info.email)
    if info.phone:
        contact_parts.append(info.phone)
    if info.location:
        contact_parts.append(info.location)
    if info.linkedin:
        contact_parts.append(info.linkedin)
    if info.portfolio:
        contact_parts.append(info.portfolio)

    if contact_parts:
        p = doc.add_paragraph()
        run = p.add_run("  ·  ".join(contact_parts))
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # ── Professional Summary ───────────────────────────────────────────────
    if cv_data.professional_summary:
        add_spacer()
        add_heading("Profile", level=2)
        p = doc.add_paragraph(cv_data.professional_summary)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # ── Work Experience ────────────────────────────────────────────────────
    if cv_data.work_experience:
        add_spacer()
        add_heading("Experience", level=2)
        for job in cv_data.work_experience:
            dates = ""
            if job.start_date:
                dates = job.start_date
                dates += f" – {'Present' if job.is_current else (job.end_date or '')}"
            add_two_col(job.job_title, dates, left_bold=True)
            company_str = job.company
            if job.location:
                company_str += f"  ·  {job.location}"
            add_sub(company_str)
            for point in job.bullet_points:
                add_bullet(point)
            add_spacer()

    # ── Education ──────────────────────────────────────────────────────────
    if cv_data.education:
        add_heading("Education", level=2)
        for edu in cv_data.education:
            add_two_col(edu.degree, edu.graduation_year or "", left_bold=True)
            add_sub(edu.institution)
            meta = []
            if edu.location:
                meta.append(edu.location)
            if edu.honors:
                meta.append(edu.honors)
            if meta:
                add_body("  ·  ".join(meta))
            add_spacer()

    # ── Skills ─────────────────────────────────────────────────────────────
    if cv_data.skills:
        add_heading("Skills", level=2)
        for group_name, items in cv_data.skills.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            label = p.add_run(f"{group_name}: ")
            label.bold = True
            label.font.size = Pt(9.5)
            label.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            value = p.add_run(", ".join(items))
            value.font.size = Pt(9.5)
            value.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        add_spacer()

    # ── Projects ───────────────────────────────────────────────────────────
    if cv_data.projects:
        add_heading("Projects", level=2)
        for proj in cv_data.projects:
            title_line = proj.title
            if proj.link:
                title_line += f"  —  {proj.link}"
            add_two_col(title_line, "", left_bold=True)
            add_body(proj.description)
            if proj.tech_stack:
                p = doc.add_paragraph()
                run = p.add_run("Stack: " + ", ".join(proj.tech_stack))
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
                run.bold = True
                p.paragraph_format.space_after = Pt(3)

    # ── Certifications ──────────────────────────────────────────────────────
    if cv_data.certifications:
        add_heading("Certifications", level=2)
        for cert in cv_data.certifications:
            # cert is a StructuredCvCertification object; build a readable label
            if hasattr(cert, "name"):
                label = cert.name
                if cert.issuer:
                    label += f" — {cert.issuer}"
                if cert.expiration_date:
                    label += f" (exp. {cert.expiration_date})"
                elif cert.issue_date:
                    label += f" ({cert.issue_date})"
            else:
                label = str(cert)
            add_bullet(label)

    # ── Serialise to bytes ─────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Cloudinary Upload ───────────────────────────────────────────────────────

def upload_to_cloudinary(
    file_bytes: bytes,
    public_id: str,
    resource_type: str = "raw",
    format: str = "pdf",
) -> dict:
    """Upload file bytes to Cloudinary and return the response dict."""
    response = cloudinary.uploader.upload(
        file_bytes,
        public_id=public_id,
        resource_type=resource_type,
        format=format,
        overwrite=True,
        access_mode="public",   # Ensure the secure_url is publicly accessible without auth
    )
    return response


# ── Orchestration Service ───────────────────────────────────────────────────

async def render_cv_documents_service(
    submission_id: str,
    payload,               # DocumentRenderRequest
    current_admin: Admin,
    db: Session,
):
    """
    Admin-triggered rendering pipeline:
    1. Verify submission & generation record exist.
    2. Deserialise StructuredCvData from stored JSON.
    3. Render requested formats (PDF / DOCX).
    4. Upload each file to Cloudinary.
    5. Save Document records with version increments.
    6. Return list of created DocumentResponse dicts.
    """
    # 1. Validate submission
    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if not submission:
        return error_response(
            status_code=status.HTTP_404_NOT_FOUND,
            message="Submission not found",
        )

    # Sub-admin RBAC: can only render for assigned submissions
    if (
        current_admin.role == AdminRole.SUB_ADMIN.value
        and submission.assigned_to_id != current_admin.id
    ):
        return error_response(
            status_code=status.HTTP_403_FORBIDDEN,
            message="You are not assigned to this submission",
        )

    # 2. Fetch generation record
    generation = db.query(AiGeneration).filter(
        AiGeneration.id == payload.ai_generation_id,
        AiGeneration.submission_id == submission_id,
    ).first()
    if not generation:
        return error_response(
            status_code=status.HTTP_404_NOT_FOUND,
            message="AI generation record not found for this submission",
        )

    # 3. Recover StructuredCvData from the generation record
    #    (stored as JSON in ai_generation.structured_cv_json)
    if not hasattr(generation, "structured_cv_json") or not generation.structured_cv_json:
        return error_response(
            status_code=status.HTTP_400_BAD_REQUEST,
            message="No structured CV data found in this generation record. Please re-generate the CV first.",
        )

    try:
        cv_data = StructuredCvData.model_validate(generation.structured_cv_json)
    except Exception as e:
        return error_response(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            message=f"Stored CV data is invalid: {e}",
        )

    # 4. Render HTML once (shared for PDF)
    html = render_cv_to_html(cv_data)

    safe_name = (submission.client.first_name + "_" + submission.client.last_name).replace(" ", "_") \
        if submission.client else f"submission_{submission_id[:8]}"

    created_docs = []
    formats = [f.lower().strip() for f in payload.formats]

    for fmt in formats:
        if fmt not in ("pdf", "docx"):
            continue

        # Determine next version
        existing_version = (
            db.query(Document)
            .filter(
                Document.submission_id == submission_id,
                Document.file_type == fmt,
            )
            .count()
        )
        version = existing_version + 1
        file_name = f"{safe_name}_CV_v{version}.{fmt}"
        cloudinary_public_id = f"ai_cv_generator/documents/{submission_id}/{fmt}_v{version}"

        try:
            if fmt == "pdf":
                file_bytes = render_pdf_bytes(html)
            else:
                file_bytes = render_docx_bytes(cv_data)
        except Exception as e:
            logger.error(f"Rendering {fmt.upper()} failed: {e}")
            return error_response(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                message=f"Failed to render {fmt.upper()} document: {str(e)}",
            )

        # Upload to Cloudinary
        try:
            cloud_resp = upload_to_cloudinary(
                file_bytes=file_bytes,
                public_id=cloudinary_public_id,
                resource_type="raw",
                format=fmt,
            )
            file_url = cloud_resp["secure_url"]
            cloud_public_id = cloud_resp["public_id"]
        except Exception as e:
            logger.error(f"Cloudinary upload failed: {e}")
            return error_response(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                message=f"File upload failed: {str(e)}",
            )

        # Save document record
        doc = Document(
            submission_id=submission_id,
            ai_generation_id=generation.id,
            file_url=file_url,
            file_name=file_name,
            public_id=cloud_public_id,
            file_type=fmt,
            version=version,
        )
        db.add(doc)
        db.flush()

        created_docs.append({
            "id": doc.id,
            "submission_id": doc.submission_id,
            "ai_generation_id": doc.ai_generation_id,
            "file_name": doc.file_name,
            "file_url": doc.file_url,
            "public_id": doc.public_id,
            "file_type": doc.file_type,
            "version": doc.version,
            "created_at": doc.created_at,
        })

    db.commit()

    return success_response(
        status_code=status.HTTP_201_CREATED,
        message=f"CV rendered successfully in {len(created_docs)} format(s)",
        data=created_docs,
    )


async def list_submission_documents_service(
    submission_id: str,
    current_admin: Admin,
    db: Session,
):
    """Returns all document records for a given submission."""
    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if not submission:
        return error_response(status_code=status.HTTP_404_NOT_FOUND, message="Submission not found")

    docs = (
        db.query(Document)
        .filter(Document.submission_id == submission_id)
        .order_by(Document.file_type, Document.version)
        .all()
    )

    return success_response(
        status_code=status.HTTP_200_OK,
        message="Documents retrieved successfully",
        data=[
            {
                "id": d.id,
                "submission_id": d.submission_id,
                "ai_generation_id": d.ai_generation_id,
                "file_name": d.file_name,
                "file_url": d.file_url,
                "public_id": d.public_id,
                "file_type": d.file_type,
                "version": d.version,
                "created_at": d.created_at,
            }
            for d in docs
        ],
    )


async def download_document_service(
    submission_id: str,
    document_id: str,
    db: Session,
):
    """
    Fetches document metadata from DB and returns bytes for direct download.
    Downloads the file from Cloudinary and streams it as a binary response.
    """
    import httpx

    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.submission_id == submission_id,
    ).first()
    if not doc:
        return None, error_response(
            status_code=status.HTTP_404_NOT_FOUND,
            message="Document not found",
        )

    try:
        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.get(doc.file_url)
            response.raise_for_status()
            file_bytes = response.content
    except Exception as e:
        logger.error(f"Failed to fetch document from Cloudinary: {e}")
        return None, error_response(
            status_code=status.HTTP_502_BAD_GATEWAY,
            message="Could not retrieve the document file. Please try again.",
        )

    return doc, file_bytes


async def client_download_document_service(
    submission_id: str,
    document_id: str,
    client_submission,  # The submission validated by client token
    db: Session,
):
    """Client-facing download — validates the submission belongs to the calling client."""
    if client_submission.id != submission_id:
        return None, error_response(
            status_code=status.HTTP_403_FORBIDDEN,
            message="You do not have access to this document",
        )
    return await download_document_service(submission_id, document_id, db)


async def list_client_documents_service(
    submission_id: str,
    access_token: str,
    db: Session,
):
    """Client-facing document listing protected by submission access token."""
    submission = db.query(Submission).filter(
        Submission.id == submission_id,
        Submission.access_token == access_token,
    ).first()
    if not submission:
        return error_response(
            status_code=status.HTTP_404_NOT_FOUND,
            message="Submission not found or access token is invalid",
        )

    docs = (
        db.query(Document)
        .filter(Document.submission_id == submission_id)
        .order_by(Document.file_type, Document.version)
        .all()
    )

    return success_response(
        status_code=status.HTTP_200_OK,
        message="Documents retrieved successfully",
        data=[
            {
                "id": d.id,
                "submission_id": d.submission_id,
                "ai_generation_id": d.ai_generation_id,
                "file_name": d.file_name,
                "file_url": d.file_url,
                "public_id": d.public_id,
                "file_type": d.file_type,
                "version": d.version,
                "created_at": str(d.created_at),
            }
            for d in docs
        ],
    )

